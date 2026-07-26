"""
client_list.py — Client court-date report: one section per client, listing
their upcoming court dates plus Responsible/Originating Attorney and
Responsible Staff. Ported from calendar-check's clientListGenerator.js.

Responsible/Originating Attorney and Responsible Staff used to be pulled
from the matters CSV export, on the (wrong) assumption they weren't exposed
on the Clio API's Matter object — see matter_fields.py for why that's not
true and how they're fetched live instead.
"""

from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO

import requests
from docx import Document
from docx.shared import Pt

from court_calendar.matter_fields import MATTER_SYNC_FIELDS, index_report_fields_by_last_name
from court_calendar.normalizer import normalize_party_name, party_names_match
from court_calendar.store import fetch_upcoming_events
from matter_matching import fetch_open_matters


@dataclass
class ClientEntry:
    display_name: str
    responsible_attorney: str
    originating_attorney: str
    responsible_staff: str
    dates: list[dict] = field(default_factory=list)


def _match_client_row(party: str, matter_rows: dict[str, dict]) -> dict | None:
    party_norm = normalize_party_name(party)
    if party_norm in matter_rows:
        return matter_rows[party_norm]
    for last, row in matter_rows.items():
        if party_names_match(party_norm, last):
            return row
    return None


def build_client_list(session: requests.Session) -> list[ClientEntry]:
    matters = fetch_open_matters(session, fields=MATTER_SYNC_FIELDS)
    matter_rows = index_report_fields_by_last_name(matters)
    events = fetch_upcoming_events()

    clients: dict[str, ClientEntry] = {}
    for e in events:
        row = _match_client_row(e["party"], matter_rows)
        if not row:
            continue

        display_name = (row.get("display_number") or e["party"]).strip()
        if display_name not in clients:
            clients[display_name] = ClientEntry(
                display_name=display_name,
                responsible_attorney=row.get("responsible_attorney") or "",
                originating_attorney=row.get("originating_attorney") or "",
                responsible_staff=row.get("responsible_staff") or "",
            )

        dt = datetime.fromisoformat(e["datetime"])
        clients[display_name].dates.append({
            "when": f"{dt.strftime('%b')} {dt.day}, {dt.year} at {dt.strftime('%H:%M')}",
            "dept": e["dept"],
            "purpose": e["purpose"] or e["purpose_raw"],
            "case_number": e["case_number"],
        })

    return sorted(clients.values(), key=lambda c: c.display_name)


def render_html(clients: list[ClientEntry]) -> str:
    generated = datetime.now().strftime("%B %d, %Y %H:%M")
    rows = []
    for c in clients:
        date_lines = "".join(
            f"<li>{d['when']} — Dept {d['dept']} — {d['purpose']} ({d['case_number']})</li>"
            for d in c.dates
        )
        rows.append(f"""
        <section style="margin-bottom:1.5rem;">
          <h3>{c.display_name}</h3>
          <p style="color:#666;font-size:0.9rem;">
            Responsible Attorney: {c.responsible_attorney or "—"} &nbsp;|&nbsp;
            Originating Attorney: {c.originating_attorney or "—"} &nbsp;|&nbsp;
            Responsible Staff: {c.responsible_staff or "—"}
          </p>
          <ul>{date_lines}</ul>
        </section>""")

    return f"""<!doctype html><html><head><meta charset="utf-8">
    <title>Client Court Dates — {generated}</title>
    <style>body{{font-family:Arial,sans-serif;max-width:800px;margin:2rem auto;}}</style>
    </head><body>
    <h1>Client Court Dates</h1>
    <p style="color:#666;">Generated {generated} — today and future events only</p>
    {"".join(rows)}
    </body></html>"""


def render_docx(clients: list[ClientEntry]) -> BytesIO:
    doc = Document()
    doc.add_heading("Client Court Dates", level=1)
    doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y %H:%M')} — today and future events only")

    for c in clients:
        doc.add_heading(c.display_name, level=2)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Responsible Attorney: {c.responsible_attorney or '—'}   "
            f"Originating Attorney: {c.originating_attorney or '—'}   "
            f"Responsible Staff: {c.responsible_staff or '—'}"
        ).italic = True
        for d in c.dates:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{d['when']} — Dept {d['dept']} — {d['purpose']} ({d['case_number']})").font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
